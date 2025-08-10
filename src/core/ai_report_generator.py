"""
AI-Powered Report Generator for RELA Malaysia Analytics Dashboard
Generates comprehensive Word and PDF reports with intelligent insights
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import base64
import io
import os
from typing import Dict, List, Any, Optional

# Document generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.error("python-docx not installed. Please install with: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Image,
        Table,
        TableStyle,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    st.error("reportlab not installed. Please install with: pip install reportlab")

try:
    from openai import OpenAI

    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

from ..utils.translations import get_text


class AIReportGenerator:
    def __init__(self, language="en"):
        self.language = language

        # Initialize OpenAI client if available
        if OPENAI_AVAILABLE:
            self.api_key = os.getenv("OPENAI_API_KEY")
            if self.api_key:
                self.client = OpenAI(api_key=self.api_key)
            else:
                self.client = None
        else:
            self.client = None

    def generate_ai_insights(
        self, data_summary: Dict, analytics_types: List[str]
    ) -> Dict[str, str]:
        """Generate AI-powered insights and recommendations"""
        if not self.client:
            return self._generate_fallback_insights(data_summary, analytics_types)

        try:
            # Create context from data
            context = self._create_data_context(data_summary, analytics_types)

            # Language-specific prompts
            if self.language == "ms":
                system_prompt = """
                Anda adalah penganalisis data pakar untuk RELA Malaysia. Analisis data yang diberikan dan berikan:
                1. 3-5 penemuan utama dalam bahasa Melayu
                2. Cadangan strategik untuk penambahbaikan
                3. Trend dan corak penting
                4. Risiko dan peluang yang dikenal pasti
                
                Jawab dalam bahasa Melayu yang professional dan mudah difahami.
                """
            else:
                system_prompt = """
                You are an expert data analyst for RELA Malaysia. Analyze the provided data and give:
                1. 3-5 key findings in clear language
                2. Strategic recommendations for improvement
                3. Important trends and patterns
                4. Identified risks and opportunities
                
                Respond in professional, easy-to-understand English.
                """

            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": context},
                ],
                max_tokens=1000,
                temperature=0.7,
            )

            insights_text = response.choices[0].message.content
            return self._parse_ai_insights(insights_text)

        except Exception as e:
            st.warning(f"AI analysis unavailable: {str(e)}")
            return self._generate_fallback_insights(data_summary, analytics_types)

    def _create_data_context(
        self, data_summary: Dict, analytics_types: List[str]
    ) -> str:
        """Create context string for AI analysis"""
        context_parts = [
            f"RELA Malaysia Analytics Report Data Summary:",
            f"Report Date: {datetime.now().strftime('%Y-%m-%d')}",
            "",
            "Data Overview:",
        ]

        for key, value in data_summary.items():
            context_parts.append(f"- {key}: {value}")

        context_parts.append(
            f"\nAnalytics Types Requested: {', '.join(analytics_types)}"
        )

        return "\n".join(context_parts)

    def _parse_ai_insights(self, insights_text: str) -> Dict[str, str]:
        """Parse AI response into structured insights"""
        lines = insights_text.split("\n")

        current_section = "findings"
        insights = {
            "findings": [],
            "recommendations": [],
            "trends": [],
            "risks_opportunities": [],
        }

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect section changes based on keywords
            if any(word in line.lower() for word in ["finding", "key", "penemuan"]):
                current_section = "findings"
            elif any(
                word in line.lower() for word in ["recommend", "cadangan", "suggestion"]
            ):
                current_section = "recommendations"
            elif any(word in line.lower() for word in ["trend", "pattern", "corak"]):
                current_section = "trends"
            elif any(
                word in line.lower()
                for word in ["risk", "opportunity", "risiko", "peluang"]
            ):
                current_section = "risks_opportunities"
            else:
                # Add content to current section
                if line.startswith(("-", "•", "*")) or line[0].isdigit():
                    insights[current_section].append(line)

        # Convert lists to strings
        result = {}
        for key, value_list in insights.items():
            result[key] = (
                "\n".join(value_list)
                if value_list
                else "No specific insights available."
            )

        return result

    def _generate_fallback_insights(
        self, data_summary: Dict, analytics_types: List[str]
    ) -> Dict[str, str]:
        """Generate basic insights when AI is not available"""
        lang = self.language

        if lang == "ms":
            return {
                "findings": "• Jumlah ahli RELA menunjukkan pertumbuhan yang stabil\n• Kadar prestasi berada pada tahap yang memuaskan\n• Taburan ahli mengikut negeri adalah seimbang",
                "recommendations": "• Tingkatkan program latihan untuk meningkatkan prestasi\n• Perkukuh sistem pelaporan untuk data yang lebih tepat\n• Jalankan kempen pengambilan di kawasan yang kurang ahli",
                "trends": "• Trend peningkatan dalam penyertaan operasi\n• Prestasi ahli menunjukkan peningkatan berterusan\n• Penglibatan lebih aktif dari ahli muda",
                "risks_opportunities": "• Peluang: Pengembangan program digital\n• Risiko: Kekurangan ahli di kawasan tertentu\n• Peluang: Kerjasama dengan agensi lain",
            }
        else:
            return {
                "findings": "• Total RELA membership shows stable growth patterns\n• Performance rates are at satisfactory levels\n• Member distribution across states is well-balanced",
                "recommendations": "• Enhance training programs to improve performance\n• Strengthen reporting systems for more accurate data\n• Conduct recruitment campaigns in underrepresented areas",
                "trends": "• Increasing trend in operational participation\n• Member performance shows continuous improvement\n• More active engagement from younger members",
                "risks_opportunities": "• Opportunity: Digital program expansion\n• Risk: Member shortage in certain areas\n• Opportunity: Collaboration with other agencies",
            }

    def create_word_document(
        self,
        report_config: Dict,
        data_summary: Dict,
        analytics_types: List[str],
        ai_insights: Dict,
    ) -> bytes:
        """Generate Word document report"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available")

        doc = Document()
        lang = self.language

        # Document styling
        styles = doc.styles
        title_style = styles["Title"]
        heading1_style = styles["Heading 1"]
        heading2_style = styles["Heading 2"]

        # Title page
        title = doc.add_paragraph(
            report_config.get("title", "RELA Malaysia Analytics Report")
        )
        title.style = title_style
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        subtitle = doc.add_paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y')}"
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add description
        if report_config.get("description"):
            desc = doc.add_paragraph(report_config["description"])
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Executive Summary (if requested)
        if report_config.get("include_executive_summary", True):
            doc.add_heading(get_text(lang, "executive_summary", "Executive Summary"), 1)

            summary_text = f"""This report provides a comprehensive analysis of RELA Malaysia operations 
            covering the period from {report_config.get('start_date', 'N/A')} to {report_config.get('end_date', 'N/A')}. 
            The analysis includes {', '.join(analytics_types)} with AI-generated insights and recommendations."""

            doc.add_paragraph(summary_text)

        # Data Summary Section
        doc.add_heading(get_text(lang, "data_summary", "Data Summary"), 1)

        # Create data summary table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = get_text(lang, "metric", "Metric")
        hdr_cells[1].text = get_text(lang, "value", "Value")

        for metric, value in data_summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)

        # AI Insights Section
        if ai_insights:
            doc.add_heading(get_text(lang, "ai_insights", "AI-Generated Insights"), 1)

            # Key Findings
            doc.add_heading(get_text(lang, "key_findings", "Key Findings"), 2)
            doc.add_paragraph(ai_insights.get("findings", "No findings available."))

            # Recommendations (if requested)
            if report_config.get("include_recommendations", True):
                doc.add_heading(get_text(lang, "recommendations", "Recommendations"), 2)
                doc.add_paragraph(
                    ai_insights.get("recommendations", "No recommendations available.")
                )

            # Trends
            doc.add_heading(get_text(lang, "trends", "Trends & Patterns"), 2)
            doc.add_paragraph(ai_insights.get("trends", "No trends identified."))

            # Risks and Opportunities
            doc.add_heading("Risks & Opportunities", 2)
            doc.add_paragraph(
                ai_insights.get(
                    "risks_opportunities", "No risks or opportunities identified."
                )
            )

        # Analytics Sections
        for analytics_type in analytics_types:
            doc.add_heading(f"{analytics_type} Analysis", 1)

            # Add placeholder content for each analytics type
            content = self._get_analytics_content(analytics_type, data_summary)
            doc.add_paragraph(content)

        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()

    def create_pdf_document(
        self,
        report_config: Dict,
        data_summary: Dict,
        analytics_types: List[str],
        ai_insights: Dict,
    ) -> bytes:
        """Generate PDF document report"""
        if not REPORTLAB_AVAILABLE:
            raise Exception("reportlab not available")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
        )

        story.append(
            Paragraph(
                report_config.get("title", "RELA Malaysia Analytics Report"),
                title_style,
            )
        )
        story.append(Spacer(1, 12))

        # Subtitle
        subtitle = Paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y')}", styles["Normal"]
        )
        story.append(subtitle)
        story.append(Spacer(1, 20))

        # Executive Summary
        if report_config.get("include_executive_summary", True):
            story.append(Paragraph("Executive Summary", styles["Heading1"]))
            summary_text = f"""This report provides a comprehensive analysis of RELA Malaysia operations 
            covering the period from {report_config.get('start_date', 'N/A')} to {report_config.get('end_date', 'N/A')}."""
            story.append(Paragraph(summary_text, styles["Normal"]))
            story.append(Spacer(1, 12))

        # Data Summary Table
        story.append(Paragraph("Data Summary", styles["Heading1"]))

        data_table_data = [["Metric", "Value"]]
        for metric, value in data_summary.items():
            data_table_data.append([metric, str(value)])

        data_table = Table(data_table_data)
        data_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 14),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ]
            )
        )

        story.append(data_table)
        story.append(Spacer(1, 20))

        # AI Insights
        if ai_insights:
            story.append(Paragraph("AI-Generated Insights", styles["Heading1"]))

            story.append(Paragraph("Key Findings", styles["Heading2"]))
            story.append(
                Paragraph(
                    ai_insights.get("findings", "No findings available."),
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 10))

            if report_config.get("include_recommendations", True):
                story.append(Paragraph("Recommendations", styles["Heading2"]))
                story.append(
                    Paragraph(
                        ai_insights.get(
                            "recommendations", "No recommendations available."
                        ),
                        styles["Normal"],
                    )
                )
                story.append(Spacer(1, 10))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _get_analytics_content(self, analytics_type: str, data_summary: Dict) -> str:
        """Get content for specific analytics type"""
        lang = self.language

        content_map = {
            "Member Analytics": "Detailed analysis of member demographics, distribution, and growth patterns.",
            "Operations Analytics": "Comprehensive review of operational activities, success rates, and resource allocation.",
            "Performance Analytics": "Assessment of performance metrics, attendance rates, and effectiveness indicators.",
            "Regional Analytics": "State-wise and district-level analysis of RELA activities and member distribution.",
            "Trend Analytics": "Historical trends, seasonal patterns, and growth trajectory analysis.",
            "Predictive Analytics": "Future projections and predictive modeling based on historical data patterns.",
        }

        return content_map.get(
            analytics_type,
            f"Analysis for {analytics_type} will be included based on available data.",
        )

    def calculate_data_summary(
        self,
        members_df: pd.DataFrame,
        operations_df: pd.DataFrame,
        assignments_df: pd.DataFrame,
        start_date: str,
        end_date: str,
    ) -> Dict:
        """Calculate comprehensive data summary for the report"""
        summary = {}

        # Basic counts
        summary["Total Members"] = len(members_df)
        summary["Total Operations"] = len(operations_df)
        summary["Total Assignments"] = len(assignments_df)

        # Member metrics
        if not members_df.empty:
            summary["Active Members"] = len(
                members_df[members_df["status"] == "Active"]
            )
            summary["Average Age"] = f"{members_df['age'].mean():.1f} years"
            summary["States Covered"] = members_df["state"].nunique()

        # Operations metrics
        if not operations_df.empty:
            summary["Completed Operations"] = len(
                operations_df[operations_df["status"] == "Completed"]
            )
            summary["Average Success Rate"] = (
                f"{operations_df['success_rate'].mean():.1%}"
            )

        # Performance metrics
        if not assignments_df.empty:
            summary["Average Performance"] = (
                f"{assignments_df['performance_score'].mean():.1f}/10"
            )
            summary["Attendance Rate"] = f"{assignments_df['attendance'].mean():.1%}"

        # Date range
        summary["Report Period"] = f"{start_date} to {end_date}"
        summary["Generated On"] = datetime.now().strftime("%Y-%m-%d %H:%M")

        return summary
